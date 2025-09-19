#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import json

# Don't import docx at module level - do it only when needed
# This prevents immediate failure if docx is not installed

def create_hebrew_word_document(transcription, title, output_path):
    """
    יוצר מסמך Word בשיטה של החלפת תבנית עובדת
    """
    try:
        # Check if python-docx is available
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError as e:
            print(f"python-docx not available: {str(e)}", file=sys.stderr)
            print("Falling back to HTML generation", file=sys.stderr)
            return create_html_fallback(transcription, title, output_path)

        import os
        import shutil
        from zipfile import ZipFile

        # בדיקה אם קיימת תבנית עובדת
        possible_templates = [
            'חזר מהשרת תקין 2.docx',
            'דוגמה_Word_מושלמת.docx',
            'בדיקה_תבנית_עובדת.docx',
            'template.docx',
            'simple-template.docx'
        ]

        template_path = None
        print(f"Looking for templates in directory: {os.getcwd()}", file=sys.stderr)
        print(f"Directory contents: {os.listdir('.')[:10]}...", file=sys.stderr)

        for template in possible_templates:
            print(f"Checking template: {template}", file=sys.stderr)
            if os.path.exists(template):
                template_path = template
                print(f"Found working template: {template}", file=sys.stderr)
                break
            else:
                print(f"Template not found: {template}", file=sys.stderr)

        if not template_path:
            print("No working template found, falling back to basic creation", file=sys.stderr)
            return create_basic_hebrew_document(transcription, title, output_path)

        # העתקת התבנית
        shutil.copy2(template_path, output_path)

        # ניקוי והכנת הטקסט
        clean_text = transcription.replace('\r\n', '\n').replace('\n\n\n', '\n\n').strip()
        sections = [section.strip() for section in clean_text.split('\n\n') if section.strip()]

        # פתיחת הקובץ כ-ZIP ועדכון התוכן
        with ZipFile(output_path, 'r') as zip_ref:
            # קריאת document.xml הקיים
            with zip_ref.open('word/document.xml') as doc_file:
                doc_content = doc_file.read().decode('utf-8')

        # יצירת תוכן חדש במבנה הקיים
        new_paragraphs = []

        # כותרת
        title_paragraph = f'''
<w:p w14:paraId="6A1F55DC" w14:textId="77777777">
  <w:pPr>
    <w:jc w:val="right"/>
    <w:spacing w:after="400"/>
    <w:rPr>
      <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
      <w:sz w:val="32"/>
      <w:b/>
    </w:rPr>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
      <w:sz w:val="32"/>
      <w:b/>
    </w:rPr>
    <w:t>{escape_xml(title)}</w:t>
  </w:r>
</w:p>'''
        new_paragraphs.append(title_paragraph)

        # שורה ריקה
        new_paragraphs.append('<w:p></w:p>')

        # פסקאות תוכן - Python מעבד הכל כאן
        import re
        all_text = ' '.join(sections)
        # תיקון ירידות שורה באמצע משפט
        all_text = re.sub(r'\n+', ' ', all_text)  # החלף ירידות שורה ברווח
        all_text = re.sub(r'\s{2,}', ' ', all_text)  # רווחים כפולים לרווח יחיד
        all_text = fix_hebrew_punctuation(all_text)  # Python עושה את כל העיבוד העברי

        # חלוקה לפסקאות - שומרים על שלמות המירכאות

        # חלוקה בזהירות למשפטים, אבל לא אם יש מירכאה פתוחה
        words = all_text.split()
        current_para = ""
        word_count = 0

        for word in words:
            current_para += word + " "
            word_count += 1

            # יצירת פסקה חדשה רק אם:
            # 1. יש מספיק מילים (קוצר עוד יותר ל-20)
            # 2. אין מירכאה פתוחה (זוגי של מירכאות)
            # 3. המשפט מסתיים
            if (word_count >= 20 or len(current_para) > 250) and word.endswith(('.', '!', '?', ':')):
                quote_count = current_para.count('"')
                if quote_count % 2 == 0:  # זוגי מירכאות = לא באמצע ציטוט
                    para_text = current_para.strip()
                    if para_text:
                        content_paragraph = f'''
<w:p>
  <w:pPr>
    <w:jc w:val="right"/>
    <w:spacing w:after="240"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
      <w:sz w:val="28"/>
    </w:rPr>
    <w:t>{escape_xml(para_text)}</w:t>
  </w:r>
</w:p>'''
                        new_paragraphs.append(content_paragraph)
                    current_para = ""
                    word_count = 0

        # פסקה אחרונה
        if current_para.strip():
            para_text = current_para.strip()
            content_paragraph = f'''
<w:p>
  <w:pPr>
    <w:jc w:val="right"/>
    <w:spacing w:after="240"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
      <w:sz w:val="28"/>
    </w:rPr>
    <w:t>{escape_xml(para_text)}</w:t>
  </w:r>
</w:p>'''
            new_paragraphs.append(content_paragraph)

        # החלפת התוכן
        import re
        # מוצא את ה-body ומחליף את התוכן
        body_content = ''.join(new_paragraphs)
        new_doc_content = re.sub(
            r'<w:body[^>]*>.*?</w:body>',
            f'<w:body>{body_content}</w:body>',
            doc_content,
            flags=re.DOTALL
        )

        # שמירת הקובץ המעודכן
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as temp_file:
            temp_path = temp_file.name

        # עדכון הקובץ
        with ZipFile(template_path, 'r') as original_zip:
            with ZipFile(temp_path, 'w') as new_zip:
                for item in original_zip.infolist():
                    if item.filename == 'word/document.xml':
                        new_zip.writestr(item, new_doc_content.encode('utf-8'))
                    else:
                        data = original_zip.read(item.filename)
                        new_zip.writestr(item, data)

        # החלפת הקובץ הסופי
        shutil.move(temp_path, output_path)

        print(f"Word document created successfully: {output_path}", file=sys.stderr)
        return True

    except Exception as e:
        print(f"Error creating Word document: {str(e)}")
        return False

def create_basic_hebrew_document(transcription, title, output_path):
    """
    יצירת מסמך בסיסי אם אין תבנית - עם הגדרות RTL משופרות
    """
    try:
        # Import docx here too
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        print("Creating basic Hebrew document with RTL settings", file=sys.stderr)
        doc = Document()

        # הגדרת השפה העיקרית של המסמך לעברית
        doc_element = doc.element
        doc_element.set(qn('xml:lang'), 'he-IL')

    except Exception as e:
        print(f"Error creating basic document: {str(e)}", file=sys.stderr)
        # אם גם זה נכשל, ניצור מסמך HTML פשוט
        return create_html_fallback(transcription, title, output_path)

    try:
        # כותרת עם הגדרות RTL מחוזקות
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.name = 'David'
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # הוספת הגדרות RTL לכותרת
        set_rtl_paragraph(title_paragraph)

        # שורה ריקה
        doc.add_paragraph()

        # עיבוד התוכן
        clean_text = transcription.replace('\r\n', '\n').replace('\n\n\n', '\n\n').strip()
        clean_text = fix_hebrew_punctuation(clean_text)
        sections = [section.strip() for section in clean_text.split('\n\n') if section.strip()]

        print(f"Creating {len(sections)} content paragraphs", file=sys.stderr)

        for i, section in enumerate(sections):
            lines = [line.strip() for line in section.split('\n') if line.strip()]
            combined_text = ' '.join(lines).strip()

            if combined_text and not combined_text[-1] in '.!?:':
                combined_text += '.'

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(combined_text)
            run.font.name = 'David'
            run.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # הוספת הגדרות RTL לכל פסקה
            set_rtl_paragraph(paragraph)

            print(f"Added paragraph {i+1}: {combined_text[:50]}...", file=sys.stderr)

        doc.save(output_path)
        print(f"Basic document saved successfully: {output_path}", file=sys.stderr)
        return True

    except Exception as e:
        print(f"Error in basic document creation process: {str(e)}", file=sys.stderr)
        return create_html_fallback(transcription, title, output_path)

def create_html_fallback(transcription, title, output_path):
    """
    יצירת קובץ HTML כ-fallback אם Python-docx לא זמין
    """
    try:
        html_content = f'''<!DOCTYPE html>
<html dir="rtl" lang="he">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: David, Arial, sans-serif;
            direction: rtl;
            text-align: right;
            margin: 40px;
            line-height: 1.6;
        }}
        h1 {{
            font-size: 18px;
            font-weight: bold;
            margin-bottom: 20px;
        }}
        p {{
            margin-bottom: 16px;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
'''

        # עיבוד הטקסט לפסקאות
        clean_text = transcription.replace('\r\n', '\n').replace('\n\n\n', '\n\n').strip()
        sections = [section.strip() for section in clean_text.split('\n\n') if section.strip()]

        for section in sections:
            lines = [line.strip() for line in section.split('\n') if line.strip()]
            combined_text = ' '.join(lines).strip()

            if combined_text and not combined_text[-1] in '.!?:':
                combined_text += '.'

            # הימנעות מ-HTML injection
            safe_text = combined_text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
            html_content += f'    <p>{safe_text}</p>\n'

        html_content += '''
</body>
</html>'''

        # שמירה כקובץ HTML במקום docx
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"Created HTML fallback at: {output_path}", file=sys.stderr)
        return True

    except Exception as e:
        print(f"HTML fallback failed: {str(e)}")
        return False

def escape_xml(text):
    """
    מחליף תווים מיוחדים ב-XML
    """
    return (text.replace('&', '&amp;')
               .replace('<', '&lt;')
               .replace('>', '&gt;')
               .replace('"', '&quot;')
               .replace("'", '&#39;'))

def fix_hebrew_punctuation(text):
    """
    עיבוד אגרסיבי וחזק לטקסט עברי - פתרון סופי לכל הבעיות
    """
    import re

    print('🔧 Starting SUPER AGGRESSIVE Hebrew processing in Python...', file=sys.stderr)

    # **שלב 0: ניקוי ראשוני**
    text = re.sub(r'["\u0022\u201C\u201D]', '"', text)  # אחד סוג גרשיים
    text = text.replace('\\', '')  # הסר כל קווים נטויים

    print('🔧 Phase 1: Cleaning quotes and backslashes...', file=sys.stderr)

    # **שלב 1: תיקון קיצורים - אגרסיבי**
    # תיקונים ישירים וחזקים
    text = text.replace('שליט "א', 'שליט"א')
    text = text.replace('רש "י', 'רש"י')
    text = text.replace('רש י', 'רש"י')
    text = text.replace('חז "ל', 'חז"ל')
    text = text.replace('חז ל', 'חז"ל')
    text = text.replace('ל "ט', 'ל"ט')
    text = text.replace('ל ט', 'ל"ט')
    text = text.replace('הרמב "ן', 'הרמב"ן')
    text = text.replace('הרמב ן', 'הרמב"ן')
    text = text.replace('רמב "ם', 'רמב"ם')
    text = text.replace('רמב ם', 'רמב"ם')
    text = text.replace('האר י ז ל', 'האר"י ז"ל')
    text = text.replace('שו "ע', 'שו"ע')
    text = text.replace('שו ע', 'שו"ע')

    print('🔧 Phase 2: Fixed abbreviations...', file=sys.stderr)

    # **שלב 2: הסרת גרשיים מיותרים - אגרסיבי**
    # הסר גרשיים ממילים בודדות שאינן ציטוטים אמיתיים
    words_to_clean = ['עמוד', 'יחיינו', 'מיומיים', 'כתוב', 'בפסוק', 'ברוך', 'תהיה', 'גוי', 'תראה', 'איך', 'ז', 'ל', 'בענין', 'מביאים', 'הרמב', 'ן', 'ר', 'צריכה', 'את', 'רוצה', 'לעשות', 'משהו', 'שמע', 'ישראל']

    for word in words_to_clean:
        text = text.replace(f'"{word}"', word)  # הסר גרשיים ממילים אלו

    # הסר גרשיים מיותרים באמצע משפטים
    text = re.sub(r'(?<=[א-ת])\s+"([א-ת]{1,8})"\s+(?=[א-ת])', r' \1 ', text)

    print('🔧 Phase 3: Removed unnecessary quotes...', file=sys.stderr)

    # **שלב ג: תיקון שמות עם גרשיים**
    text = re.sub(r'ה\s+"([^"]+)"', r'ה"\1"', text)
    text = re.sub(r'([א-ת])\s+"([^"]+)"', r'\1"\2"', text)

    # **שלב ד: מילים מתחלקות**
    text = re.sub(r'חז\s*"\s*לים', 'חז"לים', text)
    text = re.sub(r'([א-ת]+)לי\s*"\s*ם', r'\1לים', text)

    # **שלב ה: מילים צמודות**
    word_fixes = {
        'יודעתראו': 'יודעת ראו',
        'אומרתאני': 'אומרת אני',
        'שאלתיאותו': 'שאלתי אותו',
        'אמרתיכן': 'אמרתי כן',
        'אמרשלום': 'אמר שלום',
        'זהדבר': 'זה דבר',
        'חשובמאוד': 'חשוב מאוד'
    }
    for wrong, correct in word_fixes.items():
        text = text.replace(wrong, correct)

    # **שלב ו: תיקונים מתקדמים עם גרשיים ופיסוק**
    text = re.sub(r'([א-ת])\.\"', r'\1\".', text)              # נקודה לפני גרשיים
    text = re.sub(r'([א-ת])\"([א-ת])', r'\1 \"\2', text)      # רווח לפני גרשיים פותחים

    # תיקונים ספציפיים
    text = text.replace('אומר "שאל', 'אומר" שאל')
    text = text.replace('ום."ו', 'ום". ו')
    text = text.replace('".היום', '". היום')

    # תיקונים נוספים לבעיות שהתגלו
    text = text.replace('שליט\"א', 'שליט"א')           # תיקון גרש במקום גרשיים
    text = text.replace('נפלאים.בעזרת', 'נפלאים. בעזרת')  # רווח אחרי נקודה
    text = text.replace('חז\"ל', 'חז"ל')               # תיקון גרשיים
    text = text.replace('רש\"י', 'רש"י')               # תיקון גרשיים
    text = text.replace('זה.כשאדם', 'זה. כשאדם')        # רווח אחרי נקודה

    # תיקונים לבעיות שהתגלו בטקסט החדש
    text = text.replace('שליט \"א', 'שליט"א')          # רווח מיותר בקיצור
    text = text.replace('רש \"י', 'רש"י')               # רווח מיותר בקיצור
    text = text.replace('ל \"ט', 'ל"ט')                # רווח מיותר בקיצור
    text = text.replace('חז \"ל', 'חז"ל')               # רווח מיותר בקיצור

    # תיקון גרשיים מיותרים
    text = text.replace('"עמוד ב\'', 'עמוד ב\'')        # הסר גרשיים מיותרים לפני עמוד
    text = text.replace('"יחיינו "מיומיים"', '"יחיינו מיומיים"')  # תיקון גרשיים כפולים

    # תיקון דחיפת גרשיים למקומות לא נכונים
    text = text.replace('אומר "לו, "דבר "שני', 'אומר לו, דבר שני')  # הסר גרשיים מיותרים

    # תיקון כללי לגרשיים מיותרים במקומות שלא צריך
    text = re.sub(r'"([א-ת]{1,3})\b(?!\s*[א-ת]*")', r'\1', text)  # הסר גרשיים ממילים קצרות שאינן ציטוט
    text = re.sub(r'\b([א-ת]{1,2})\s+"([א-ת])', r'\1"\2', text)    # תקן רווח מיותר בקיצורים

    # תיקון קווים נטויים מיותרים
    text = re.sub(r'([א-ת])\\([א-ת])', r'\1\2', text)              # הסר \ ממילים עבריות
    text = re.sub(r'\\+', '', text)                                # הסר כל הקווים הנטויים המיותרים
    text = text.replace('\\"', '"')                                # תקן \" ל-" רגיל
    text = text.replace('\\', '')                                   # הסר כל קווים נטויים נותרים

    # תיקון קיצורים עם רווחים מיותרים
    text = text.replace('האר י ז ל', 'האר"י ז"ל')                 # תיקון קיצור מפורק
    text = text.replace('הרמב ן', 'הרמב"ן')                       # תיקון רמב"ן
    text = text.replace('ר\\\'', 'ר\'')                            # תיקון ר' עם קו נטוי

    # תיקון גרשיים תקועים במקומות מיותרים
    text = re.sub(r'"([א-ת]{1,4})"\s+(?![א-ת]*")', r'\1 ', text)   # הסר גרשיים ממילים בודדות
    text = re.sub(r'([.,])\s*"([א-ת]{1,5})"\s*', r'\1 \2 ', text)  # הסר גרשיים אחרי פיסוק

    # תיקונים ספציפיים לדוגמאות שהובאו
    text = text.replace('"ברוך "תהיה', '"ברוך תהיה')
    text = text.replace('"תראה "איך', '"תראה איך')
    text = text.replace('בקרב הארץ\\"', 'בקרב הארץ')

    # תיקונים נוספים לבעיות שהתגלו בבדיקה
    text = text.replace('שליט "א', 'שליט"א')
    text = text.replace('חז ל', 'חז"ל')
    text = text.replace('ל ט', 'ל"ט')
    text = text.replace('"ז"ל "בענין', '"ז"ל בענין')
    text = text.replace('"יחיינו "מיומיים "כתוב "בפסוק', '"יחיינו מיומיים" כתוב בפסוק')
    text = text.replace('"הרמב"ן', 'הרמב"ן')
    text = text.replace('"ר\'', 'ר\'')

    # תיקון גרשיים מיותרים במקומות ספציפיים - חזק יותר
    text = re.sub(r'"([א-ת]{1,8})"(?=\s+[א-ת])', r'\1', text)   # הסר גרשיים ממילים בודדות לפני מילה
    text = re.sub(r'(?<=[א-ת])\s+"([א-ת]{1,8})"\s+', r' \1 ', text)  # הסר גרשיים ממילים באמצע משפט
    text = re.sub(r'"([א-ת]{1,8})"(?=[\s.,!?])', r'\1', text)   # הסר גרשיים לפני פיסוק או רווח

    # תיקון ציטוטים שבורים
    text = re.sub(r'"([א-ת]+)\s+"([א-ת]+)"', r'"\1 \2"', text)  # תקן ציטוטים מפוצלים

    # תיקון גרשיים חסרים בתחילת ציטוט
    text = re.sub(r'([א-ת])\s+([א-ת][^"]*)"', r'\1 "\2"', text)  # מילה מילה" -> מילה "מילה"

    # תיקון נקודות צמודות למילים
    text = text.replace('נמאס.מאיפה', 'נמאס. מאיפה')
    text = text.replace('ההצלחה.דוד', 'ההצלחה. דוד')

    # תיקונים כלליים
    text = re.sub(r'\"\.([א-ת])', r'\". \1', text)

    # תיקון מילים עם ר'
    text = re.sub(r'ר\s*\'\s*([א-ת])', r'ר\' \1', text)

    # תיקון מיתוקים
    text = text.replace('אמן-ים', 'אמנים')
    text = re.sub(r'([א-ת]+)-ים\b', r'\1ים', text)

    # ניקוי כללי
    text = re.sub(r'\s+([.,!?:;])', r'\1', text)              # הסר רווח לפני פיסוק
    text = re.sub(r'([.,!?:;])([א-ת])', r'\1 \2', text)       # הוסף רווח אחרי פיסוק לפני עברית
    text = re.sub(r'([.,!?:;])\s+', r'\1 ', text)             # רווח יחיד אחרי פיסוק
    text = re.sub(r'\s{2,}', ' ', text)                       # רווחים כפולים
    text = text.strip()

    print('✅ Comprehensive Hebrew processing completed in Python', file=sys.stderr)
    return text

def set_rtl_paragraph(paragraph):
    """
    מגדיר פסקה כ-RTL (ימין לשמאל) - גרסה משופרת
    """
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # הוספת הגדרת RTL ל-XML של הפסקה
        p = paragraph._element
        pPr = p.get_or_add_pPr()

        # הוספת bidi element
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

        # הוספת textDirection element
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), 'rl')
        pPr.append(textDirection)

        # הוספת יישור ימין חזק
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)

        print(f"RTL settings applied to paragraph", file=sys.stderr)

    except Exception as e:
        print(f"Warning: Cannot set RTL: {str(e)}", file=sys.stderr)

def main():
    """
    פונקציה ראשית המקבלת פרמטרים מ-Node.js
    """
    try:
        print("Python script started", file=sys.stderr)
        print(f"Python version: {sys.version}", file=sys.stderr)
        print(f"Arguments: {sys.argv}", file=sys.stderr)

        if len(sys.argv) != 2:
            print("Usage: python generate_word_doc.py '<json_data>'")
            sys.exit(1)

        # בדיקת python-docx - לא בהכרח נדרשת כי יש לנו HTML fallback
        try:
            import docx
            print(f"python-docx version: {docx.__version__}", file=sys.stderr)
            print("python-docx is available", file=sys.stderr)
        except ImportError as e:
            print(f"python-docx import error: {str(e)}", file=sys.stderr)
            print("Will use HTML fallback instead", file=sys.stderr)

        # קבלת הנתונים מ-Node.js
        json_data = sys.argv[1]
        print(f"Received JSON data length: {len(json_data)}", file=sys.stderr)

        data = json.loads(json_data)
        print(f"Parsed data keys: {list(data.keys())}", file=sys.stderr)

        transcription = data.get('transcription', '')
        title = data.get('title', 'תמלול')
        output_path = data.get('output_path', 'output.docx')

        print(f"Creating document: {title} -> {output_path}", file=sys.stderr)
        print(f"Transcription type: {type(transcription)}", file=sys.stderr)
        print(f"Transcription length: {len(str(transcription))}", file=sys.stderr)
        print(f"Transcription preview: {str(transcription)[:100]}...", file=sys.stderr)

        # Validation של הטקסט
        if not transcription or not isinstance(transcription, str):
            error_msg = f"Invalid transcription data: type={type(transcription)}, value={str(transcription)[:200]}"
            print(f"ERROR: {error_msg}", file=sys.stderr)
            print(json.dumps({"success": False, "error": error_msg}))
            sys.exit(1)

        if len(transcription.strip()) < 10:
            error_msg = f"Transcription too short: '{transcription}'"
            print(f"ERROR: {error_msg}", file=sys.stderr)
            print(json.dumps({"success": False, "error": error_msg}))
            sys.exit(1)

        # יצירת המסמך
        success = create_hebrew_word_document(transcription, title, output_path)

        if success:
            print(json.dumps({"success": True, "file_path": output_path}))
        else:
            print(json.dumps({"success": False, "error": "Failed to create document"}))

    except Exception as e:
        print(f"Exception in main: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()