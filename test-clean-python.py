#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def main():
    test_text = """שלום, זה בדיקת תמלול מהשרת החדש עם פתרון Python!

הטקסט הזה מכיל פסקאות מרובות. האם הוא יעבוד כהלכה עם הפתרון החדש? אני מקווה שכן.

זו פסקה נוספת עם סימני פיסוק: נקודות, פסיקים, סימני קריאה! וגם נקודותיים: כמו כאן.

פסקה רביעית לבדיקת פרדת הפסקאות. זה אמור להיראות מעולה ללא בעיות פיסוק.

בואו נראה איך זה עובד בפועל!"""

    doc = Document()

    # כותרת
    title = doc.add_paragraph()
    title_run = title.add_run("בדיקה סופית של פתרון Python")
    title_run.font.name = 'David'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # הגדרת RTL לכותרת
    title_pPr = title._element.get_or_add_pPr()
    title_bidi = OxmlElement('w:bidi')
    title_bidi.set(qn('w:val'), '1')
    title_pPr.append(title_bidi)

    # שורה ריקה
    doc.add_paragraph()

    # עיבוד פסקאות
    sections = [s.strip() for s in test_text.split('\n\n') if s.strip()]

    for section in sections:
        lines = [line.strip() for line in section.split('\n') if line.strip()]
        combined_text = ' '.join(lines).strip()

        # הוספת נקודה בסוף אם אין
        if combined_text and not combined_text[-1] in '.!?:':
            combined_text += '.'

        # יצירת פסקה
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(combined_text)
        run.font.name = 'David'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # הגדרת RTL
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    # שמירה
    doc.save('test_clean.docx')
    print("File created: test_clean.docx")

if __name__ == "__main__":
    main()