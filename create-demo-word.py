#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_demo_word():
    """
    יוצר קובץ דוגמה Word עם תוכן עברי מושלם
    """
    doc = Document()

    # כותרת ראשית
    title = doc.add_paragraph()
    title_run = title.add_run("דוגמה לקובץ Word עברי מושלם")
    title_run.font.name = 'David'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # הגדרת RTL לכותרת
    title_pPr = title._element.get_or_add_pPr()
    title_bidi = OxmlElement('w:bidi')
    title_bidi.set(qn('w:val'), '1')
    title_pPr.append(title_bidi)

    # שורה ריקה
    doc.add_paragraph()

    # תוכן הדוגמה
    demo_text = """שלום וברכה! זהו קובץ דוגמה שנוצר עם הפתרון Python החדש.

הטקסט הזה מדגים את כל היכולות של המערכת החדשה. האם אתה רואה שהכל עובד כהלכה? כן, זה עובד מצוין!

פסקה זו מכילה סימני פיסוק מגוונים: נקודות, פסיקים, סימני קריאה! וגם נקודותיים: בדיוק כמו שצריך.

הפתרון החדש מבוסס Python ופותר את כל הבעיות שהיו במערכת הישנה. אין עוד טקסט דחוס, אין בעיות יישור, והכי חשוב - אין סימני פיסוק קופצים!

זו הפסקה האחרונה בדוגמה. אני מקווה שאתה מרוצה מהתוצאה הסופית."""

    # עיבוד פסקאות
    sections = [s.strip() for s in demo_text.split('\n\n') if s.strip()]

    for section in sections:
        lines = [line.strip() for line in section.split('\n') if line.strip()]
        combined_text = ' '.join(lines).strip()

        # הוספת נקודה בסוף אם אין
        if combined_text and not combined_text[-1] in '.!?:':
            combined_text += '.'

        # יצירת פסקה
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(combined_text)
        run.font.name = 'David'
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # הגדרת RTL
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    # הוספת שורה ריקה
    doc.add_paragraph()

    # הוספת חתימה
    signature = doc.add_paragraph()
    sig_run = signature.add_run("נוצר עם הפתרון Python המשופר ✓")
    sig_run.font.name = 'David'
    sig_run.font.size = Pt(10)
    sig_run.italic = True
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # הגדרת RTL לחתימה
    sig_pPr = signature._element.get_or_add_pPr()
    sig_bidi = OxmlElement('w:bidi')
    sig_bidi.set(qn('w:val'), '1')
    sig_pPr.append(sig_bidi)

    # שמירה
    output_file = 'Demo_Hebrew_Perfect.docx'
    doc.save(output_file)
    print(f"Demo Word file created successfully: {output_file}")

    return output_file

if __name__ == "__main__":
    create_demo_word()