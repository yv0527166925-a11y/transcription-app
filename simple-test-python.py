#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_simple_hebrew_doc():
    doc = Document()

    # כותרת
    title = doc.add_paragraph()
    title_run = title.add_run("בדיקה פשוטה")
    title_run.font.name = 'David'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # הגדרת RTL לכותרת
    title_pPr = title._element.get_or_add_pPr()
    title_bidi = OxmlElement('w:bidi')
    title_bidi.set(qn('w:val'), '1')
    title_pPr.append(title_bidi)

    # תוכן
    content = doc.add_paragraph()
    content_run = content.add_run("זה טקסט פשוט לבדיקה. האם הוא עובד כהלכה? כן, נראה טוב!")
    content_run.font.name = 'David'
    content_run.font.size = Pt(12)
    content.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # הגדרת RTL לתוכן
    content_pPr = content._element.get_or_add_pPr()
    content_bidi = OxmlElement('w:bidi')
    content_bidi.set(qn('w:val'), '1')
    content_pPr.append(content_bidi)

    # שמירה
    doc.save('בדיקה_פשוטה.docx')
    print("File created: בדיקה_פשוטה.docx")

if __name__ == "__main__":
    create_simple_hebrew_doc()