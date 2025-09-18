#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_fixed_rtl_doc():
    """
    יוצר מסמך Word עם RTL חזק ופיסוק תקין
    """
    doc = Document()

    # הגדרת RTL ברמת המסמך
    doc_element = doc.element.body
    sectPr = doc_element.get_or_add_sectPr()
    bidi_doc = OxmlElement('w:bidi')
    sectPr.append(bidi_doc)

    # כותרת
    title = doc.add_paragraph()
    title_run = title.add_run("בדיקה עם RTL חזק ופיסוק תקין")
    title_run.font.name = 'David'
    title_run.font.size = Pt(16)
    title_run.bold = True

    # יישור וRTL חזק לכותרת
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_pPr = title._element.get_or_add_pPr()

    # הגדרות RTL מרובות לוודא שזה עובד
    title_jc = OxmlElement('w:jc')
    title_jc.set(qn('w:val'), 'right')
    title_pPr.append(title_jc)

    title_bidi = OxmlElement('w:bidi')
    title_bidi.set(qn('w:val'), '1')
    title_pPr.append(title_bidi)

    title_textDirection = OxmlElement('w:textDirection')
    title_textDirection.set(qn('w:val'), 'rl')
    title_pPr.append(title_textDirection)

    # RTL ברמת ה-run
    title_rPr = title_run._element.get_or_add_rPr()
    title_rtl = OxmlElement('w:rtl')
    title_rPr.append(title_rtl)

    # הגדרת שפה עברית
    title_lang = OxmlElement('w:lang')
    title_lang.set(qn('w:val'), 'he-IL')
    title_lang.set(qn('w:eastAsia'), 'he-IL')
    title_lang.set(qn('w:bidi'), 'he-IL')
    title_rPr.append(title_lang)

    # שורה ריקה
    doc.add_paragraph()

    # תוכן מבחן
    test_content = [
        "זהו טקסט מבחן ראשון. האם הוא מיושר לימין כהלכה?",
        "פסקה שנייה עם פיסוק: נקודות, פסיקים, סימני קריאה! האם זה עובד?",
        "פסקה שלישית לבדיקה סופית. הכל אמור להיראות מושלם עכשיו."
    ]

    for content_text in test_content:
        # יצירת פסקה עם הגדרות RTL מרובות
        paragraph = doc.add_paragraph()

        # הגדרת יישור ראשונה
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # הגדרות XML מפורטות
        pPr = paragraph._element.get_or_add_pPr()

        # יישור מפורש
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'right')
        pPr.append(jc)

        # bidi
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

        # כיוון טקסט
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), 'rl')
        pPr.append(textDirection)

        # יצירת ה-run
        run = paragraph.add_run(content_text)
        run.font.name = 'David'
        run.font.size = Pt(12)

        # הגדרות RTL ברמת ה-run
        rPr = run._element.get_or_add_rPr()

        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)

        # שפה עברית
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), 'he-IL')
        lang.set(qn('w:eastAsia'), 'he-IL')
        lang.set(qn('w:bidi'), 'he-IL')
        rPr.append(lang)

        # פונט מפורש
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'David')
        rFonts.set(qn('w:hAnsi'), 'David')
        rFonts.set(qn('w:cs'), 'David')
        rPr.append(rFonts)

    # שמירה
    output_file = 'Fixed_RTL_Test.docx'
    doc.save(output_file)
    print(f"Fixed RTL document created: {output_file}")

    return output_file

if __name__ == "__main__":
    create_fixed_rtl_doc()